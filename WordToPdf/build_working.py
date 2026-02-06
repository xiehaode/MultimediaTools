"""
可工作的构建脚本 - 绕过Cython问题，直接创建可用的解决方案
"""

import os
import sys
import shutil
from pathlib import Path

def create_simple_ctypes_wrapper():
    """创建简单的ctypes包装器来模拟DLL"""
    print("创建ctypes包装器...")
    
    wrapper_code = '''
"""
WordToPdf ctypes包装器 - 模拟C DLL行为
"""

import ctypes
import os
import sys
from pathlib import Path

class MockWordToPdfDLL:
    """模拟WordToPdf DLL的行为"""
    
    def __init__(self):
        self._last_error = ""
        
    def convert_file(self, input_path, output_path):
        """模拟 convert_file 函数"""
        try:
            # 导入实际的转换函数
            from WordToPdf import convert_docx_to_pdf, convert_pdf_to_docx
            
            self._last_error = ""
            
            if not os.path.exists(input_path):
                self._last_error = f"输入文件不存在: {input_path}"
                return 0
            
            input_ext = Path(input_path).suffix.lower()
            output_ext = Path(output_path).suffix.lower()
            
            if input_ext == '.docx' and output_ext == '.pdf':
                convert_docx_to_pdf(input_path, output_path)
                return 1
            elif input_ext == '.pdf' and output_ext == '.docx':
                convert_pdf_to_docx(input_path, output_path)
                return 1
            else:
                self._last_error = f"不支持的格式转换: {input_ext} -> {output_ext}"
                return 0
                
        except ImportError:
            self._last_error = "WordToPdf模块未找到"
            return 0
        except Exception as e:
            self._last_error = str(e)
            return 0
    
    def convert_files_batch(self, input_paths, file_count, output_dir, output_format):
        """模拟批量转换函数"""
        try:
            self._last_error = ""
            
            # 确保输出目录存在
            os.makedirs(output_dir, exist_ok=True)
            
            success_count = 0
            for i in range(file_count):
                input_path = input_paths[i].decode('utf-8') if isinstance(input_paths[i], bytes) else input_paths[i]
                
                if not os.path.exists(input_path):
                    continue
                
                # 生成输出文件名
                input_name = Path(input_path).stem
                output_path = os.path.join(output_dir, f"{input_name}.{output_format}")
                
                if self.convert_file(input_path, output_path):
                    success_count += 1
            
            return 1 if success_count > 0 else 0
            
        except Exception as e:
            self._last_error = str(e)
            return 0
    
    def check_file_support(self, file_path):
        """检查文件支持"""
        try:
            self._last_error = ""
            
            if not os.path.exists(file_path):
                self._last_error = f"文件不存在: {file_path}"
                return 0
            
            file_ext = Path(file_path).suffix.lower()
            supported = ['.pdf', '.docx']
            return 1 if file_ext in supported else 0
            
        except Exception as e:
            self._last_error = str(e)
            return 0
    
    def get_supported_formats(self):
        """获取支持的格式"""
        self._last_error = ""
        return '{"pdf": "docx", "docx": "pdf"}'
    
    def get_file_type(self, file_path):
        """获取文件类型"""
        try:
            self._last_error = ""
            
            if not os.path.exists(file_path):
                self._last_error = f"文件不存在: {file_path}"
                return "error"
            
            file_ext = Path(file_path).suffix.lower()
            type_map = {
                '.pdf': 'pdf',
                '.docx': 'docx',
                '.doc': 'doc'
            }
            return type_map.get(file_ext, 'unknown')
            
        except Exception as e:
            self._last_error = str(e)
            return "error"
    
    def get_last_error(self):
        """获取最后错误"""
        return self._last_error.encode('utf-8')
    
    def clear_error(self):
        """清除错误"""
        self._last_error = ""

# 全局模拟DLL实例
_mock_dll = MockWordToPdfDLL()

# C函数接口
def convert_file(input_path, output_path):
    return _mock_dll.convert_file(input_path, output_path)

def convert_files_batch(input_paths, file_count, output_dir, output_format):
    return _mock_dll.convert_files_batch(input_paths, file_count, output_dir, output_format)

def check_file_support(file_path):
    return _mock_dll.check_file_support(file_path)

def get_supported_formats():
    return _mock_dll.get_supported_formats()

def get_file_type(file_path):
    return _mock_dll.get_file_type(file_path)

def get_last_error():
    return _mock_dll.get_last_error()

def clear_error():
    _mock_dll.clear_error()
'''
    
    with open("word_converter_ctypes.py", "w", encoding="utf-8") as f:
        f.write(wrapper_code)
    
    print("创建了 word_converter_ctypes.py")
    return True

def create_simple_test():
    """创建简单测试脚本"""
    test_code = '''
"""
WordToPdf 转换测试脚本
"""

import sys
import os
from pathlib import Path

# 导入转换器
try:
    from word_converter_ctypes import convert_file, get_last_error
    print("使用ctypes包装器")
except ImportError:
    try:
        from word_converter_native import convert_file, get_last_error
        print("使用native包装器")
    except ImportError:
        print("错误: 无法导入任何转换器包装器")
        sys.exit(1)

def test_conversion():
    """测试转换功能"""
    print("WordToPdf 转换测试")
    print("=" * 40)
    
    # 检查是否提供了参数
    if len(sys.argv) >= 3:
        input_file = sys.argv[1]
        output_file = sys.argv[2]
    else:
        # 使用默认测试文件
        input_file = "test.docx"
        output_file = "test_output.pdf"
        
        # 如果测试文件不存在，创建一个简单的
        if not os.path.exists(input_file):
            print(f"创建测试文件: {input_file}")
            try:
                from docx import Document
                doc = Document()
                doc.add_heading("测试文档", 0)
                doc.add_paragraph("这是一个测试文档，用于验证WordToPdf转换功能。")
                doc.add_paragraph("如果看到这个PDF，说明转换成功了！")
                doc.save(input_file)
                print(f"测试文件创建成功: {input_file}")
            except ImportError:
                print("无法创建测试文件，请提供输入文件参数")
                return
    
    # 检查输入文件
    if not os.path.exists(input_file):
        print(f"错误: 输入文件不存在 - {input_file}")
        return
    
    print(f"输入文件: {input_file}")
    print(f"输出文件: {output_file}")
    
    # 执行转换
    print("\n开始转换...")
    try:
        result = convert_file(input_file, output_file)
        
        if result:
            print("✅ 转换成功!")
            if os.path.exists(output_file):
                print(f"输出文件: {os.path.abspath(output_file)}")
                print(f"文件大小: {os.path.getsize(output_file)} 字节")
            else:
                print("警告: 转换成功但输出文件未找到")
        else:
            error = get_last_error()
            if isinstance(error, bytes):
                error = error.decode('utf-8')
            print(f"❌ 转换失败: {error}")
            
    except Exception as e:
        print(f"❌ 转换过程中发生错误: {e}")

def show_supported_formats():
    """显示支持的格式"""
    try:
        if 'word_converter_ctypes' in sys.modules:
            from word_converter_ctypes import get_supported_formats, get_file_type
        else:
            from word_converter_native import get_supported_formats, get_file_type
            
        print("\n支持的格式:")
        print(get_supported_formats())
        
        # 测试文件类型检测
        test_files = ["test.docx", "test.pdf", "test.txt"]
        print("\n文件类型检测:")
        for file in test_files:
            file_type = get_file_type(file)
            print(f"  {file} -> {file_type}")
            
    except Exception as e:
        print(f"无法获取格式信息: {e}")

if __name__ == "__main__":
    # 显示支持格式
    show_supported_formats()
    
    # 执行转换测试
    test_conversion()
    
    print("\n使用方法:")
    print("1. 不带参数 - 使用默认测试文件")
    print("2. 带参数 - python test_converter.py input.docx output.pdf")
'''
    
    with open("test_converter.py", "w", encoding="utf-8") as f:
        f.write(test_code)
    
    print("创建了 test_converter.py")
    return True

def create_requirements():
    """创建requirements文件"""
    requirements = [
        "docx2pdf>=0.1.8",
        "pdf2docx>=0.5.0", 
        "python-docx>=0.8.11",
        "reportlab>=3.6.0",
        "Pillow>=8.0.0"
    ]
    
    with open("requirements_simple.txt", "w") as f:
        f.write("\n".join(requirements))
    
    print("创建了 requirements_simple.txt")
    return True

def main():
    """主函数"""
    print("WordToPdf 可工作解决方案构建器")
    print("=" * 50)
    
    # 创建包装器
    create_simple_ctypes_wrapper()
    
    # 创建测试脚本
    create_simple_test()
    
    # 创建需求文件
    create_requirements()
    
    print("\n构建完成!")
    print("\n创建的文件:")
    print("  - word_converter_ctypes.py    # ctypes包装器")
    print("  - test_converter.py         # 测试脚本")
    print("  - requirements_simple.txt    # 依赖列表")
    
    print("\n使用方法:")
    print("1. 安装依赖: pip install -r requirements_simple.txt")
    print("2. 运行测试: python test_converter.py")
    print("3. 转换文件: python test_converter.py input.docx output.pdf")
    
    print("\n提示:")
    print("- 这个方案绕过了Cython编译问题")
    print("- 提供了与C DLL相同的API接口")
    print("- 可以直接在Python项目中使用")
    print("- 也可以作为C++ DLL的原型")

if __name__ == "__main__":
    main()