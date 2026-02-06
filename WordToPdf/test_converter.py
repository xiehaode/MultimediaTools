"""
测试通用文件转换器
"""

import os
import tempfile
import pandas as pd
from PIL import Image, ImageDraw
from docx import Document
from universal_converter import UniversalConverter

def create_test_files():
    """创建测试文件"""
    test_dir = "test_files"
    os.makedirs(test_dir, exist_ok=True)
    
    # 创建测试Word文档
    doc = Document()
    doc.add_heading('测试文档', 0)
    doc.add_paragraph('这是一个测试段落，用于测试Word转PDF功能。')
    doc.add_paragraph('第二段内容，包含中文字符测试。')
    
    # 添加表格
    table = doc.add_table(rows=3, cols=3)
    for i in range(3):
        for j in range(3):
            table.cell(i, j).text = f'单元格{i+1}-{j+1}'
    
    doc_path = os.path.join(test_dir, 'test.docx')
    doc.save(doc_path)
    
    # 创建测试CSV文件
    data = {
        '姓名': ['张三', '李四', '王五'],
        '年龄': [25, 30, 35],
        '城市': ['北京', '上海', '广州']
    }
    df = pd.DataFrame(data)
    csv_path = os.path.join(test_dir, 'test.csv')
    df.to_csv(csv_path, index=False, encoding='utf-8-sig')
    
    # 创建测试图片
    img = Image.new('RGB', (400, 300), color='white')
    draw = ImageDraw.Draw(img)
    draw.rectangle([50, 50, 350, 250], outline='black', width=2)
    draw.text((100, 150), "测试图片", fill='black')
    
    img_path = os.path.join(test_dir, 'test.png')
    img.save(img_path)
    
    # 创建测试HTML文件
    html_content = """
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>测试页面</title>
    <style>
        body { font-family: Arial, sans-serif; margin: 20px; }
        h1 { color: #333; }
        table { border-collapse: collapse; width: 100%; }
        th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
    </style>
</head>
<body>
    <h1>测试HTML文档</h1>
    <p>这是一个测试HTML文档，用于测试HTML转PDF功能。</p>
    <table>
        <tr><th>姓名</th><th>年龄</th><th>城市</th></tr>
        <tr><td>张三</td><td>25</td><td>北京</td></tr>
        <tr><td>李四</td><td>30</td><td>上海</td></tr>
        <tr><td>王五</td><td>35</td><td>广州</td></tr>
    </table>
</body>
</html>
"""
    html_path = os.path.join(test_dir, 'test.html')
    with open(html_path, 'w', encoding='utf-8') as f:
        f.write(html_content)
    
    return {
        'word': doc_path,
        'csv': csv_path,
        'image': img_path,
        'html': html_path
    }

def test_converter():
    """测试转换器功能"""
    print("=== 通用文件转换器测试 ===\n")
    
    # 创建转换器实例
    converter = UniversalConverter()
    
    # 显示支持的转换格式
    print("支持的转换格式:")
    conversions = converter.get_supported_conversions()
    for from_type, to_types in conversions.items():
        print(f"  {from_type.upper()} -> {', '.join(to_types)}")
    print()
    
    # 创建测试文件
    print("创建测试文件...")
    test_files = create_test_files()
    print("测试文件创建完成\n")
    
    # 测试各种转换
    test_cases = [
        ('word', 'test.docx', 'test_from_word.pdf'),
        ('csv', 'test.csv', 'test_from_csv.html'),
        ('image', 'test.png', 'test_from_image.pdf'),
        ('html', 'test.html', 'test_from_html.pdf'),
    ]
    
    results = []
    
    for file_type, input_file, output_file in test_cases:
        print(f"测试 {file_type.upper()} 转换...")
        
        input_path = os.path.join('test_files', input_file)
        output_path = os.path.join('test_files', output_file)
        
        if os.path.exists(input_path):
            success = converter.convert(input_path, output_path)
            
            if success and os.path.exists(output_path):
                file_size = os.path.getsize(output_path)
                results.append((file_type, True, file_size))
                print(f"  ✓ {input_file} -> {output_file} (大小: {file_size} 字节)")
            else:
                results.append((file_type, False, 0))
                print(f"  ✗ {input_file} -> {output_file} (转换失败)")
        else:
            results.append((file_type, False, 0))
            print(f"  ✗ 输入文件不存在: {input_path}")
    
    print("\n=== 测试结果汇总 ===")
    success_count = sum(1 for _, success, _ in results if success)
    total_count = len(results)
    
    print(f"总测试数: {total_count}")
    print(f"成功: {success_count}")
    print(f"失败: {total_count - success_count}")
    print(f"成功率: {success_count/total_count*100:.1f}%")
    
    for file_type, success, size in results:
        status = "✓" if success else "✗"
        print(f"  {status} {file_type.upper()}: {'成功' if success else '失败'}")
    
    # 测试批量转换
    print("\n=== 测试批量转换 ===")
    input_files = [
        os.path.join('test_files', 'test.docx'),
        os.path.join('test_files', 'test.csv'),
    ]
    output_dir = os.path.join('test_files', 'batch_output')
    
    batch_results = converter.convert_batch(input_files, output_dir, 'pdf')
    
    batch_success = sum(1 for success in batch_results.values() if success)
    print(f"批量转换完成: {batch_success}/{len(batch_results)} 成功")
    
    return success_count == total_count

if __name__ == "__main__":
    success = test_converter()
    if success:
        print("\n🎉 所有测试通过!")
    else:
        print("\n⚠️  部分测试失败，请检查错误信息")