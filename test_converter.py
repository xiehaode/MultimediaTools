
"""
WordToPdf 转换测试脚本
"""

import sys
import os
from pathlib import Path

# 导入转换器
try:
    from word_converter_ctypes import convert_file, get_last_error
    print("使用ctypes包装器")
except ImportError:
    try:
        from word_converter_native import convert_file, get_last_error
        print("使用native包装器")
    except ImportError:
        print("错误: 无法导入任何转换器包装器")
        sys.exit(1)

def test_conversion():
    """测试转换功能"""
    print("WordToPdf 转换测试")
    print("=" * 40)
    
    # 检查是否提供了参数
    if len(sys.argv) >= 3:
        input_file = sys.argv[1]
        output_file = sys.argv[2]
    else:
        # 使用默认测试文件
        input_file = "test.docx"
        output_file = "test_output.pdf"
        
        # 如果测试文件不存在，创建一个简单的
        if not os.path.exists(input_file):
            print(f"创建测试文件: {input_file}")
            try:
                from docx import Document
                doc = Document()
                doc.add_heading("测试文档", 0)
                doc.add_paragraph("这是一个测试文档，用于验证WordToPdf转换功能。")
                doc.add_paragraph("如果看到这个PDF，说明转换成功了！")
                doc.save(input_file)
                print(f"测试文件创建成功: {input_file}")
            except ImportError:
                print("无法创建测试文件，请提供输入文件参数")
                return
    
    # 检查输入文件
    if not os.path.exists(input_file):
        print(f"错误: 输入文件不存在 - {input_file}")
        return
    
    print(f"输入文件: {input_file}")
    print(f"输出文件: {output_file}")
    
    # 执行转换
    print("
开始转换...")
    try:
        result = convert_file(input_file, output_file)
        
        if result:
            print("✅ 转换成功!")
            if os.path.exists(output_file):
                print(f"输出文件: {os.path.abspath(output_file)}")
                print(f"文件大小: {os.path.getsize(output_file)} 字节")
            else:
                print("警告: 转换成功但输出文件未找到")
        else:
            error = get_last_error()
            if isinstance(error, bytes):
                error = error.decode('utf-8')
            print(f"❌ 转换失败: {error}")
            
    except Exception as e:
        print(f"❌ 转换过程中发生错误: {e}")

def show_supported_formats():
    """显示支持的格式"""
    try:
        if 'word_converter_ctypes' in sys.modules:
            from word_converter_ctypes import get_supported_formats, get_file_type
        else:
            from word_converter_native import get_supported_formats, get_file_type
            
        print("
支持的格式:")
        print(get_supported_formats())
        
        # 测试文件类型检测
        test_files = ["test.docx", "test.pdf", "test.txt"]
        print("
文件类型检测:")
        for file in test_files:
            file_type = get_file_type(file)
            print(f"  {file} -> {file_type}")
            
    except Exception as e:
        print(f"无法获取格式信息: {e}")

if __name__ == "__main__":
    # 显示支持格式
    show_supported_formats()
    
    # 执行转换测试
    test_conversion()
    
    print("
使用方法:")
    print("1. 不带参数 - 使用默认测试文件")
    print("2. 带参数 - python test_converter.py input.docx output.pdf")
